import customtkinter as ctk
import tkinter as tk
from tkinter import messagebox, filedialog, simpledialog
import pandas as pd
from datetime import datetime, timedelta
from PIL import Image, ImageFont, ImageDraw
import os
import shutil
from barcode.writer import ImageWriter
import barcode
import webbrowser
import json
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from docx import Document

class StokYonetimProgrami(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Stok Takip & Yönetim Programı v9.0")
        self.geometry("1400x800")
        
        self.current_theme = "light"
        ctk.set_appearance_mode(self.current_theme)
        ctk.set_default_color_theme("blue")
        
        self.urunler_db_path = "urunler.csv"
        self.satislar_db_path = "satislar.csv"
        self.stok_log_db_path = "stok_log.csv"
        self.borclar_db_path = "borclar.csv"
        self.borc_odemeler_db_path = "borc_odemeler.csv"
        self.musteriler_db_path = "musteriler.csv"
        self.firmalar_db_path = "firmalar.csv"
        
        self.df_urunler = self.db_yukle(self.urunler_db_path)
        self.df_satislar = self.db_yukle(self.satislar_db_path)
        self.df_stok_log = self.db_yukle(self.stok_log_db_path)
        self.df_borclar = self.db_yukle(self.borclar_db_path)
        self.df_borc_odemeler = self.db_yukle(self.borc_odemeler_db_path)
        self.df_musteriler = self.db_yukle(self.musteriler_db_path)
        self.df_firmalar = self.db_yukle(self.firmalar_db_path)

        self.saticilar = ["Ahmet", "Efe", "Orhan", "Kerem"]
        self.odeme_yontemleri = ["Nakit", "Kart", "Borç", "Taksitli"]
        self.sepet = {}
        self.borc_sepet = {}
        self.current_satis_id = None
        self.history = []
        
        self.create_widgets()
        self.bind("<F1>", self.go_back)

    def db_yukle(self, path):
        try:
            if not os.path.exists(path) or os.stat(path).st_size == 0:
                if path == self.urunler_db_path:
                    return pd.DataFrame(columns=['urun_adi', 'barkod', 'firma_adi', 'kdv', 'alis_fiyati', 'satis_fiyati', 'stok_miktari', 'ozellikler'])
                elif path == self.satislar_db_path:
                    return pd.DataFrame(columns=['satis_id', 'tarih', 'toplam_tutar', 'indirim_miktari_tl', 'toplam_kar', 'urunler', 'odeme_yontemi', 'satıcı', 'musteri_id'])
                elif path == self.stok_log_db_path:
                    return pd.DataFrame(columns=['tarih', 'barkod', 'miktar_degisimi', 'aciklama'])
                elif path == self.borclar_db_path:
                    return pd.DataFrame(columns=['satis_id', 'musteri_id', 'tarih', 'borc_miktari', 'odenmis_miktar', 'taksit_miktari', 'taksit_gunu', 'urunler'])
                elif path == self.borc_odemeler_db_path:
                    return pd.DataFrame(columns=['odeme_id', 'satis_id', 'tarih', 'miktar', 'aciklama'])
                elif path == self.musteriler_db_path:
                    return pd.DataFrame(columns=['musteri_id', 'isim', 'soyisim', 'telefon_no', 'adres'])
                elif path == self.firmalar_db_path:
                    return pd.DataFrame(columns=['firma_adi', 'adres', 'telefon', 'vergi_dairesi', 'yetkili_kisi', 'notlar'])
            
            df = pd.read_csv(path)
            if 'tarih' in df.columns: 
                df['tarih'] = pd.to_datetime(df['tarih'], errors='coerce')
            if 'barkod' in df.columns: 
                df['barkod'] = df['barkod'].astype(str)
            return df
        except pd.errors.EmptyDataError:
            return pd.DataFrame(columns=self.db_yukle(path).columns)
        except FileNotFoundError:
            return pd.DataFrame(columns=self.db_yukle(path).columns)

    def db_kaydet(self, df, path):
        df.to_csv(path, index=False)

    def go_back(self, event=None):
        if len(self.history) > 1:
            self.history.pop()
            last_screen_func = self.history.pop()
            last_screen_func()

    def clear_sag_frame(self):
        for widget in self.sag_frame.winfo_children():
            widget.destroy()

    def toggle_theme(self):
        if self.current_theme == "light":
            self.current_theme = "dark"
            ctk.set_appearance_mode("dark")
        else:
            self.current_theme = "light"
            ctk.set_appearance_mode("light")

    def create_widgets(self):
        self.sol_frame = ctk.CTkFrame(self, width=200, corner_radius=0)
        self.sol_frame.pack(side="left", fill="y")
        self.sag_frame = ctk.CTkFrame(self, corner_radius=0)
        self.sag_frame.pack(side="right", fill="both", expand=True)
        self.create_sol_menu()
        self.create_main_menu()

    def create_sol_menu(self):
        ctk.CTkButton(self.sol_frame, text="Ana Menü", command=self.create_main_menu).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sol_frame, text="Satış Yap", command=self.create_satis_form).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sol_frame, text="Ürün Yönetimi", command=self.create_urun_yonetim_ekrani).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sol_frame, text="Stok Durum Raporu", command=self.create_stok_raporu).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sol_frame, text="Satış Geçmişi", command=self.create_satis_gecmisi_ekrani).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sol_frame, text="Borç & Taksit Yönetimi", command=self.create_borc_yonetim_ekrani).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sol_frame, text="Firmalar", command=self.create_firma_yonetimi).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sol_frame, text="Yönetici Paneli", command=self.create_yonetici_ekrani_password).pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(self.sol_frame, text="Uygulamadan Çık", command=self.quit_app).pack(pady=10, padx=20, fill="x", side="bottom")

    def create_main_menu(self):
        self.clear_sag_frame()
        self.history.append(self.create_main_menu)
        
        header_frame = ctk.CTkFrame(self.sag_frame, fg_color="transparent")
        header_frame.pack(pady=10, fill="x")
        ctk.CTkLabel(header_frame, text="Ana Menü", font=("Arial", 30, "bold")).pack(side="left", padx=20)
        ctk.CTkButton(header_frame, text="Siyah/Beyaz Tema", command=self.toggle_theme).pack(side="right", padx=20)

        stok_sayisi = self.df_urunler['stok_miktari'].sum() if not self.df_urunler.empty else 0
        dusuk_stok_sayisi = len(self.df_urunler[self.df_urunler['stok_miktari'] < 5]) if not self.df_urunler.empty else 0
        musteri_sayisi = len(self.df_musteriler) if not self.df_musteriler.empty else 0
        
        cards_frame = ctk.CTkFrame(self.sag_frame, fg_color="transparent")
        cards_frame.pack(pady=10, padx=20)
        
        card_stok = ctk.CTkFrame(cards_frame, corner_radius=10)
        card_stok.pack(side="left", padx=10, pady=10)
        ctk.CTkLabel(card_stok, text="Toplam Stok", font=("Arial", 18, "bold")).pack(pady=5, padx=20)
        ctk.CTkLabel(card_stok, text=f"{int(stok_sayisi)}", font=("Arial", 40, "bold")).pack(pady=10, padx=20)
        
        card_dusuk_stok = ctk.CTkFrame(cards_frame, corner_radius=10)
        card_dusuk_stok.pack(side="left", padx=10, pady=10)
        ctk.CTkLabel(card_dusuk_stok, text="Düşük Stok", font=("Arial", 18, "bold")).pack(pady=5, padx=20)
        ctk.CTkLabel(card_dusuk_stok, text=f"{dusuk_stok_sayisi}", font=("Arial", 40, "bold"), text_color="red" if dusuk_stok_sayisi > 0 else "green").pack(pady=10, padx=20)
        
        card_musteri = ctk.CTkFrame(cards_frame, corner_radius=10)
        card_musteri.pack(side="left", padx=10, pady=10)
        ctk.CTkLabel(card_musteri, text="Kayıtlı Müşteri", font=("Arial", 18, "bold")).pack(pady=5, padx=20)
        ctk.CTkLabel(card_musteri, text=f"{musteri_sayisi}", font=("Arial", 40, "bold")).pack(pady=10, padx=20)

        yillik_ciro = self.df_satislar[pd.to_datetime(self.df_satislar['tarih']).dt.year == datetime.now().year]['toplam_tutar'].sum() if not self.df_satislar.empty else 0
        card_ciro = ctk.CTkFrame(cards_frame, corner_radius=10)
        card_ciro.pack(side="left", padx=10, pady=10)
        ctk.CTkLabel(card_ciro, text="Yıllık Ciro", font=("Arial", 18, "bold")).pack(pady=5, padx=20)
        ctk.CTkLabel(card_ciro, text=f"{yillik_ciro:.2f} TL", font=("Arial", 40, "bold")).pack(pady=10, padx=20)
        
        ctk.CTkLabel(self.sag_frame, text="Son 7 Günlük Satışlar", font=("Arial", 20, "bold")).pack(pady=10)
        self.create_satis_grafigi()

    def create_satis_grafigi(self):
        fig_frame = ctk.CTkFrame(self.sag_frame)
        fig_frame.pack(pady=10)
        
        bugun = datetime.now().date()
        son_7_gun = [(bugun - timedelta(days=i)) for i in range(7)]
        
        gunluk_satislar = {}
        for gun in son_7_gun:
            gun_str = gun.strftime('%Y-%m-%d')
            gunluk_satislar[gun_str] = self.df_satislar[self.df_satislar['tarih'].dt.date.astype(str) == gun_str]['toplam_tutar'].sum()

        gunler = [gun.strftime('%d-%m') for gun in son_7_gun]
        cirolar = [gunluk_satislar[gun.strftime('%Y-%m-%d')] for gun in son_7_gun]
        
        fig, ax = plt.subplots(figsize=(8, 4), facecolor='#ebebeb')
        fig.set_facecolor('#ebebeb')
        ax.set_facecolor('#ebebeb')
        ax.tick_params(colors='black')
        ax.spines['bottom'].set_color('black')
        ax.spines['left'].set_color('black')
        ax.set_title("Haftalık Satış Grafiği", color='black')
        ax.set_xlabel("Tarih", color='black')
        ax.set_ylabel("Ciro (TL)", color='black')

        ax.bar(gunler, cirolar, color='#3498db')
        
        canvas = FigureCanvasTkAgg(fig, master=fig_frame)
        canvas_widget = canvas.get_tk_widget()
        canvas_widget.pack(fill=tk.BOTH, expand=True)

    def create_urun_yonetim_ekrani(self):
        self.clear_sag_frame()
        self.history.append(self.create_urun_yonetim_ekrani)
        ctk.CTkLabel(self.sag_frame, text="Ürün Yönetimi", font=("Arial", 24, "bold")).pack(pady=10)
        top_frame = ctk.CTkFrame(self.sag_frame)
        top_frame.pack(pady=10, padx=20, fill="x")
        ctk.CTkButton(top_frame, text="Yeni Ürün Ekle", command=lambda: self.create_urun_form()).pack(side="left", padx=10)
        ctk.CTkButton(top_frame, text="Fatura ile Stok Ekle", command=self.fatura_ile_stok_girme_form).pack(side="left", padx=10)
        self.search_entry = ctk.CTkEntry(top_frame, placeholder_text="Ürün adı veya barkod ile ara...")
        self.search_entry.pack(side="left", padx=10, fill="x", expand=True)
        self.search_entry.bind("<KeyRelease>", self.urun_listesi_goster)
        self.urun_list_frame = ctk.CTkScrollableFrame(self.sag_frame, height=500)
        self.urun_list_frame.pack(pady=20, padx=20, fill="both", expand=True)
        self.urun_listesi_goster()

    def urun_listesi_goster(self, event=None):
        for widget in self.urun_list_frame.winfo_children():
            widget.destroy()
        arama_metni = self.search_entry.get().lower()
        if self.df_urunler.empty:
            ctk.CTkLabel(self.urun_list_frame, text="Stokta ürün bulunamadı.").pack(pady=10)
            return
        filtreli_df = self.df_urunler[
            self.df_urunler['urun_adi'].str.lower().str.contains(arama_metni, na=False) |
            self.df_urunler['barkod'].str.lower().str.contains(arama_metni, na=False)
        ]
        if filtreli_df.empty:
            ctk.CTkLabel(self.urun_list_frame, text="Aradığınız kritere uyan ürün bulunamadı.").pack(pady=10)
            return
        
        for _, row in filtreli_df.iterrows():
            urun_frame = ctk.CTkFrame(self.urun_list_frame)
            urun_frame.pack(pady=5, fill="x", padx=10)
            
            main_info_frame = ctk.CTkFrame(urun_frame, fg_color="transparent")
            main_info_frame.pack(fill="x", padx=10, pady=5)
            ctk.CTkLabel(main_info_frame, text=row['urun_adi'], font=("Arial", 14, "bold")).pack(side="left", padx=10)
            ctk.CTkLabel(main_info_frame, text=f"({row['barkod']})", font=("Arial", 14)).pack(side="left", padx=10)
            ctk.CTkLabel(main_info_frame, text=f"{row['satis_fiyati']:.2f} TL", font=("Arial", 14)).pack(side="left", padx=10)
            
            ozellikler_frame = ctk.CTkFrame(urun_frame, fg_color="transparent")
            ozellikler_frame.pack(fill="x", padx=10, pady=5)
            ctk.CTkLabel(ozellikler_frame, text=f"Stok: {int(row['stok_miktari'])} | Alış: {row['alis_fiyati']:.2f} TL | Özellikler: {row.get('ozellikler', '')}", font=("Arial", 12)).pack(side="left", padx=10)
            
            action_frame = ctk.CTkFrame(urun_frame, fg_color="transparent")
            action_frame.pack(fill="x", padx=10, pady=5)
            ctk.CTkButton(action_frame, text="Stok Değiştir", width=70, command=lambda r=row: self.stok_guncelle(r)).pack(side="right", padx=5)
            ctk.CTkButton(action_frame, text="Barkod Yazdır", width=70, command=lambda r=row: self.barkod_olustur_ve_kaydet(r['barkod'], r['urun_adi'], r['satis_fiyati'])).pack(side="right", padx=5)
            ctk.CTkButton(action_frame, text="Düzenle", width=70, command=lambda r=row: self.create_urun_form(edit_mode=True, product_data=r)).pack(side="right", padx=5)

    def fatura_ile_stok_girme_form(self):
        self.clear_sag_frame()
        self.history.append(self.fatura_ile_stok_girme_form)
        ctk.CTkLabel(self.sag_frame, text="Fatura ile Stok Girişi", font=("Arial", 24, "bold")).pack(pady=10)
        
        bilgilendirme = "Lütfen aşağıdaki alana her satıra bir ürünün barkodu ve miktarını yazın.\nÖrnek: barkod,miktar"
        ctk.CTkLabel(self.sag_frame, text=bilgilendirme).pack(pady=10)
        
        self.fatura_text = ctk.CTkTextbox(self.sag_frame, width=400, height=300)
        self.fatura_text.pack(pady=10)
        
        ctk.CTkButton(self.sag_frame, text="Stokları Güncelle", command=self.fatura_ile_stok_girme).pack(pady=10)

    def fatura_ile_stok_girme(self):
        satirlar = self.fatura_text.get("1.0", "end-1c").strip().split('\n')
        
        try:
            for satir in satirlar:
                bolumler = [s.strip() for s in satir.split(',')]
                if len(bolumler) >= 2:
                    barkod, miktar_str = bolumler[0], bolumler[1]
                    if not barkod or not miktar_str.isdigit(): continue
                    miktar = int(miktar_str)
                    
                    if barkod in self.df_urunler['barkod'].values:
                        self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'] += miktar
                        log_kaydi = {'tarih': datetime.now(), 'barkod': barkod, 'miktar_degisimi': miktar, 'aciklama': "Fatura ile stok girişi"}
                        log_satir = pd.DataFrame([log_kaydi])
                        self.df_stok_log = pd.concat([self.df_stok_log, log_satir], ignore_index=True)
            
            self.db_kaydet(self.df_urunler, self.urunler_db_path)
            self.db_kaydet(self.df_stok_log, self.stok_log_db_path)
            messagebox.showinfo("Başarılı", "Stoklar başarıyla güncellendi.")
            self.create_urun_yonetim_ekrani()
        except Exception as e:
            messagebox.showerror("Hata", f"İşlem sırasında bir hata oluştu: {e}")

    def create_urun_form(self, edit_mode=False, product_data=None):
        self.clear_sag_frame()
        self.history.append(self.create_urun_form)
        
        baslik = "Ürün Güncelleme" if edit_mode else "Yeni Ürün Ekle"
        ctk.CTkLabel(self.sag_frame, text=baslik, font=("Arial", 24, "bold")).pack(pady=10)
        
        form_frame = ctk.CTkFrame(self.sag_frame)
        form_frame.pack(pady=20, padx=20, fill="x")
        
        etiketler = ["Ürün Adı:", "Firma İsmi:", "KDV:", "Alış Fiyatı:", "Satış Fiyatı:", "Stok Miktarı:", "Özellikler:"]
        if not edit_mode: etiketler.insert(1, "Barkod:")
        
        self.urun_entry_vars = {}
        for i, etiket in enumerate(etiketler):
            ctk.CTkLabel(form_frame, text=etiket).grid(row=i, column=0, padx=10, pady=5, sticky="w")
            var = ctk.StringVar()
            self.urun_entry_vars[etiket] = var
            entry = ctk.CTkEntry(form_frame, textvariable=var)
            entry.grid(row=i, column=1, padx=10, pady=5, sticky="ew")
            if edit_mode and etiket != "Barkod:":
                var.set(str(product_data[etiket.replace(":", "").replace(" ", "_").lower()]))
            if edit_mode and etiket == "Barkod:":
                ctk.CTkLabel(form_frame, text=product_data['barkod']).grid(row=i, column=1, padx=10, pady=5, sticky="ew")

        ctk.CTkButton(form_frame, text="Kaydet", command=lambda: self.urun_kaydet(edit_mode, product_data)).grid(row=len(etiketler), column=0, columnspan=2, pady=10)

        if edit_mode:
            ctk.CTkLabel(self.sag_frame, text="Stok Değişim Geçmişi", font=("Arial", 18, "bold")).pack(pady=10)
            log_frame = ctk.CTkScrollableFrame(self.sag_frame, height=200)
            log_frame.pack(pady=10, padx=20, fill="both", expand=True)
            
            logs = self.df_stok_log[self.df_stok_log['barkod'] == product_data['barkod']].sort_values(by='tarih', ascending=False)
            if logs.empty:
                ctk.CTkLabel(log_frame, text="Bu ürün için stok değişim kaydı bulunamadı.").pack(pady=5)
            else:
                for _, log_row in logs.iterrows():
                    ctk.CTkLabel(log_frame, text=f"Tarih: {log_row['tarih'].strftime('%d-%m-%Y %H:%M')} | Miktar: {log_row['miktar_degisimi']} | Açıklama: {log_row['aciklama']}").pack(pady=2, fill="x")

    def urun_kaydet(self, edit_mode=False, product_data=None):
        try:
            urun_adi = self.urun_entry_vars["Ürün Adı:"].get()
            firma_adi = self.urun_entry_vars["Firma İsmi:"].get()
            kdv = float(self.urun_entry_vars["KDV:"].get())
            alis_fiyati = float(self.urun_entry_vars["Alış Fiyatı:"].get())
            satis_fiyati = float(self.urun_entry_vars["Satış Fiyatı:"].get())
            stok_miktari = int(self.urun_entry_vars["Stok Miktarı:"].get())
            ozellikler = self.urun_entry_vars["Özellikler:"].get()
            
            if not urun_adi or not firma_adi:
                messagebox.showwarning("Uyarı", "Ürün Adı ve Firma Adı boş bırakılamaz.")
                return
            
            self.firma_kaydet_veya_guncelle(firma_adi)
            
            if edit_mode:
                self.df_urunler.loc[self.df_urunler['barkod'] == product_data['barkod'], 'urun_adi'] = urun_adi
                self.df_urunler.loc[self.df_urunler['barkod'] == product_data['barkod'], 'firma_adi'] = firma_adi
                self.df_urunler.loc[self.df_urunler['barkod'] == product_data['barkod'], 'kdv'] = kdv
                self.df_urunler.loc[self.df_urunler['barkod'] == product_data['barkod'], 'alis_fiyati'] = alis_fiyati
                self.df_urunler.loc[self.df_urunler['barkod'] == product_data['barkod'], 'satis_fiyati'] = satis_fiyati
                self.df_urunler.loc[self.df_urunler['barkod'] == product_data['barkod'], 'stok_miktari'] = stok_miktari
                self.df_urunler.loc[self.df_urunler['barkod'] == product_data['barkod'], 'ozellikler'] = ozellikler
                self.db_kaydet(self.df_urunler, self.urunler_db_path)
                messagebox.showinfo("Başarılı", "Ürün başarıyla güncellendi.")
                self.create_urun_yonetim_ekrani()
            else:
                barkod = self.urun_entry_vars["Barkod:"].get().strip()
                if not barkod:
                    messagebox.showwarning("Uyarı", "Barkod alanı boş bırakılamaz.")
                    return
                if barkod in self.df_urunler['barkod'].values:
                    messagebox.showwarning("Uyarı", "Bu barkoda sahip bir ürün zaten mevcut.")
                    return
                
                yeni_urun = {'urun_adi': urun_adi, 'barkod': barkod, 'firma_adi': firma_adi, 'kdv': kdv, 'alis_fiyati': alis_fiyati, 'satis_fiyati': satis_fiyati, 'stok_miktari': stok_miktari, 'ozellikler': ozellikler}
                yeni_satir = pd.DataFrame([yeni_urun])
                self.df_urunler = pd.concat([self.df_urunler, yeni_satir], ignore_index=True)
                self.db_kaydet(self.df_urunler, self.urunler_db_path)
                messagebox.showinfo("Başarılı", f"Ürün başarıyla kaydedildi. Barkod: {barkod}")
                cevap = messagebox.askyesno("Barkod Oluştur", "Ürün için barkod etiketi oluşturulsun mu?")
                if cevap: self.barkod_olustur_ve_kaydet(barkod, urun_adi, satis_fiyati)
                self.create_urun_yonetim_ekrani()
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli sayısal değerler girin.")

    def barkod_olustur(self):
        son_barkod_str = self.df_urunler['barkod'].max() if not self.df_urunler.empty and not self.df_urunler['barkod'].isnull().all() else '0000000000000'
        son_barkod_int = int(son_barkod_str)
        yeni_barkod_int = son_barkod_int + 1
        yeni_barkod_str = f"{yeni_barkod_int:013}"
        return yeni_barkod_str

    def create_satis_form(self, satis_data=None):
        self.clear_sag_frame()
        self.history.append(self.create_satis_form)
        ctk.CTkLabel(self.sag_frame, text="Satış Yap", font=("Arial", 24, "bold")).pack(pady=10)
        
        satis_frame = ctk.CTkFrame(self.sag_frame)
        satis_frame.pack(pady=20, padx=20, fill="x")
        
        ctk.CTkLabel(satis_frame, text="Ürün Ara:", font=("Arial", 16)).pack(side="left", padx=(10,0))
        self.urun_arama_entry = ctk.CTkEntry(satis_frame, width=300)
        self.urun_arama_entry.pack(side="left", padx=10, fill="x", expand=True)
        self.urun_arama_entry.bind("<KeyRelease>", self.urun_listesi_satis)
        self.urun_arama_entry.bind("<Return>", self.urun_getir_satis) # Barkod okuyucu için
        
        self.sepet_liste_frame = ctk.CTkScrollableFrame(self.sag_frame, height=300)
        self.sepet_liste_frame.pack(pady=10, padx=20, fill="x")

        odeme_frame = ctk.CTkFrame(self.sag_frame)
        odeme_frame.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(odeme_frame, text="Ödeme Yöntemi:", font=("Arial", 14)).pack(side="left", padx=5)
        self.odeme_menu = ctk.CTkOptionMenu(odeme_frame, values=self.odeme_yontemleri)
        self.odeme_menu.pack(side="left", padx=5)
        
        ctk.CTkLabel(odeme_frame, text="Satıcı:", font=("Arial", 14)).pack(side="left", padx=5)
        self.satici_menu = ctk.CTkOptionMenu(odeme_frame, values=self.saticilar)
        self.satici_menu.pack(side="left", padx=5)
        
        self.musteri_frame = ctk.CTkFrame(self.sag_frame)
        self.musteri_frame.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(self.musteri_frame, text="Müşteri Adı Soyadı:", font=("Arial", 14)).pack(side="left", padx=5)
        self.musteri_isim_entry = ctk.CTkEntry(self.musteri_frame, placeholder_text="İsim")
        self.musteri_isim_entry.pack(side="left", padx=5, expand=True, fill="x")
        self.musteri_soyisim_entry = ctk.CTkEntry(self.musteri_frame, placeholder_text="Soyisim")
        self.musteri_soyisim_entry.pack(side="left", padx=5, expand=True, fill="x")
        ctk.CTkLabel(self.musteri_frame, text="Telefon:", font=("Arial", 14)).pack(side="left", padx=5)
        self.telefon_entry = ctk.CTkEntry(self.musteri_frame)
        self.telefon_entry.pack(side="left", padx=5, expand=True, fill="x")

        toplam_frame = ctk.CTkFrame(self.sag_frame)
        toplam_frame.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(toplam_frame, text="Toplam Tutar:", font=("Arial", 18, "bold")).pack(side="left", padx=10)
        self.toplam_tutar_label = ctk.CTkLabel(toplam_frame, text="0.00 TL", font=("Arial", 18, "bold"), text_color="green")
        self.toplam_tutar_label.pack(side="left", padx=10)
        
        ctk.CTkLabel(toplam_frame, text="Ödenecek Tutar:", font=("Arial", 18, "bold")).pack(side="left", padx=10)
        self.odenecek_tutar_entry = ctk.CTkEntry(toplam_frame, font=("Arial", 18, "bold"))
        self.odenecek_tutar_entry.pack(side="right", padx=10)
        self.odenecek_tutar_entry.insert(0, "0.00")

        ctk.CTkButton(self.sag_frame, text="Satışı Tamamla", command=self.satis_tamamla).pack(pady=10)
        
        self.urun_arama_frame = ctk.CTkScrollableFrame(self.sag_frame, height=200)
        self.urun_arama_frame.pack(pady=10, padx=20, fill="both", expand=True)
        self.urun_listesi_satis()
        
        if satis_data:
            self.yukle_satis_duzenle(satis_data)

    def yukle_satis_duzenle(self, satis_data):
        self.sepet = json.loads(satis_data['urunler'])
        self.current_satis_id = satis_data['satis_id']
        
        if not pd.isna(satis_data['musteri_id']):
            musteri = self.df_musteriler[self.df_musteriler['musteri_id'] == satis_data['musteri_id']].iloc[0]
            self.musteri_isim_entry.insert(0, musteri.get('isim', ''))
            self.musteri_soyisim_entry.insert(0, musteri.get('soyisim', ''))
            self.telefon_entry.insert(0, musteri.get('telefon_no', ''))

        self.odeme_menu.set(satis_data['odeme_yontemi'])
        self.satici_menu.set(satis_data['satıcı'])
        
        self.sepet_goster()

    def urun_listesi_satis(self, event=None):
        for widget in self.urun_arama_frame.winfo_children():
            widget.destroy()
        
        arama_metni = self.urun_arama_entry.get().lower()
        
        filtreli_df = self.df_urunler[
            self.df_urunler['urun_adi'].str.lower().str.contains(arama_metni, na=False) |
            self.df_urunler['barkod'].str.lower().str.contains(arama_metni, na=False)
        ]

        for _, row in filtreli_df.iterrows():
            urun_frame = ctk.CTkFrame(self.urun_arama_frame)
            urun_frame.pack(pady=5, fill="x", padx=10)
            
            ctk.CTkLabel(urun_frame, text=f"{row['urun_adi']} ({row['barkod']}) - {row['satis_fiyati']:.2f} TL (Stok: {int(row['stok_miktari'])})").pack(side="left", expand=True)
            ctk.CTkButton(urun_frame, text="Ekle", command=lambda r=row: self.urun_getir_satis(urun=r)).pack(side="right")
    
    def urun_getir_satis(self, event=None, urun=None):
        if not urun:
            query = self.urun_arama_entry.get().lower()
            if not query: return
            urunler = self.df_urunler[(self.df_urunler['barkod'] == query) | (self.df_urunler['urun_adi'].str.lower().str.contains(query, na=False))]
            if urunler.empty:
                messagebox.showerror("Hata", "Ürün bulunamadı.")
                return
            urun = urunler.iloc[0]

        barkod = urun['barkod']
        
        if int(urun['stok_miktari']) <= 0:
            messagebox.showwarning("Uyarı", "Bu ürün stokta kalmamıştır. Satışa devam ediliyor.")
        
        if barkod in self.sepet:
            self.sepet[barkod]['miktar'] += 1
        else:
            self.sepet[barkod] = {'urun_adi': urun['urun_adi'], 'alis_fiyati': urun['alis_fiyati'], 'satis_fiyati': urun['satis_fiyati'], 'miktar': 1, 'ozellikler': urun.get('ozellikler', '')}
        
        self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'] -= 1
        
        log_kaydi = {'tarih': datetime.now(), 'barkod': barkod, 'miktar_degisimi': -1, 'aciklama': "Satış (Sepet Ekleme)"}
        log_satir = pd.DataFrame([log_kaydi])
        self.df_stok_log = pd.concat([self.df_stok_log, log_satir], ignore_index=True)
        
        self.db_kaydet(self.df_urunler, self.urunler_db_path)
        self.db_kaydet(self.df_stok_log, self.stok_log_db_path)
        
        self.sepet_goster()
        self.urun_arama_entry.delete(0, 'end')
        self.urun_listesi_satis()

    def sepet_goster(self):
        for widget in self.sepet_liste_frame.winfo_children():
            widget.destroy()
        
        toplam_tutar = 0
        for barkod, item in self.sepet.items():
            urun_frame = ctk.CTkFrame(self.sepet_liste_frame)
            urun_frame.pack(pady=2, fill="x")
            
            ctk.CTkLabel(urun_frame, text=item['urun_adi'], font=("Arial", 14)).pack(side="left", padx=5, expand=True, fill="x")
            
            ctk.CTkLabel(urun_frame, text="Adet:", font=("Arial", 14)).pack(side="left", padx=5)
            miktar_entry = ctk.CTkEntry(urun_frame, width=50)
            miktar_entry.insert(0, str(item['miktar']))
            miktar_entry.pack(side="left", padx=5)
            miktar_entry.bind("<Return>", lambda e, b=barkod: self.sepet_miktar_revize(b, miktar_entry.get()))
            
            ctk.CTkLabel(urun_frame, text="Fiyat:", font=("Arial", 14)).pack(side="left", padx=5)
            fiyat_entry = ctk.CTkEntry(urun_frame, width=80)
            fiyat_entry.insert(0, f"{item['satis_fiyati']:.2f}")
            fiyat_entry.pack(side="left", padx=5)
            fiyat_entry.bind("<KeyRelease>", lambda e, b=barkod: self.sepet_fiyat_revize(b, fiyat_entry.get()))

            sil_btn = ctk.CTkButton(urun_frame, text="Sil", width=50, command=lambda b=barkod: self.sepetten_sil(b))
            sil_btn.pack(side="right", padx=5)
            
            toplam_tutar += item['satis_fiyati'] * item['miktar']
        
        self.toplam_tutar_label.configure(text=f"{toplam_tutar:.2f} TL")
        self.odenecek_tutar_entry.delete(0, 'end')
        self.odenecek_tutar_entry.insert(0, f"{toplam_tutar:.2f}")
    
    def sepet_miktar_revize(self, barkod, yeni_miktar_str):
        try:
            yeni_miktar = int(yeni_miktar_str)
            if yeni_miktar < 0: raise ValueError
            urun_stok = self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'].iloc[0]
            eski_miktar = self.sepet[barkod]['miktar']
            stok_degisimi = yeni_miktar - eski_miktar
            if urun_stok - stok_degisimi < 0:
                messagebox.showwarning("Uyarı", "Stok yeterli değil.")
                self.sepet_goster()
                return
            self.sepet[barkod]['miktar'] = yeni_miktar
            self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'] -= stok_degisimi
            self.db_kaydet(self.df_urunler, self.urunler_db_path)
            self.sepet_goster()
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli bir sayı girin.")
            self.sepet_goster()
    
    def sepet_fiyat_revize(self, barkod, yeni_fiyat_str):
        try:
            yeni_fiyat = float(yeni_fiyat_str)
            if yeni_fiyat < 0: raise ValueError
            self.sepet[barkod]['satis_fiyati'] = yeni_fiyat
            self.sepet_goster()
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli bir fiyat girin.")
            self.sepet_goster()

    def sepetten_sil(self, barkod):
        if barkod in self.sepet:
            miktar_geri = self.sepet[barkod]['miktar']
            self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'] += miktar_geri
            del self.sepet[barkod]
            self.db_kaydet(self.df_urunler, self.urunler_db_path)
            self.sepet_goster()

    def satis_tamamla(self):
        if not self.sepet:
            messagebox.showwarning("Uyarı", "Sepette ürün bulunmamaktadır.")
            return
        
        try:
            toplam_tutar = sum(item['satis_fiyati'] * item['miktar'] for item in self.sepet.values())
            odenecek_tutar = float(self.odenecek_tutar_entry.get())
            indirim_miktari = toplam_tutar - odenecek_tutar
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli bir ödenecek tutar girin.")
            return
        
        odeme_yontemi = self.odeme_menu.get()
        musteri_isim = self.musteri_isim_entry.get().strip()
        musteri_soyisim = self.musteri_soyisim_entry.get().strip()
        telefon_no = self.telefon_entry.get().strip()

        if (odeme_yontemi == "Borç" or odeme_yontemi == "Taksitli") and (not musteri_isim or not musteri_soyisim):
            messagebox.showerror("Hata", "Borç veya Taksitli satışlarda müşteri adı ve soyadı zorunludur.")
            return
        
        musteri_id = None
        if musteri_isim and musteri_soyisim:
            filtre = self.df_musteriler[(self.df_musteriler['isim'] == musteri_isim) & (self.df_musteriler['soyisim'] == musteri_soyisim)]
            if not filtre.empty:
                musteri_id = filtre.iloc[0]['musteri_id']
                if telefon_no and pd.isna(filtre.iloc[0]['telefon_no']):
                    self.df_musteriler.loc[self.df_musteriler['musteri_id'] == musteri_id, 'telefon_no'] = telefon_no
                    self.db_kaydet(self.df_musteriler, self.musteriler_db_path)
            else:
                musteri_id = self.df_musteriler['musteri_id'].max() + 1 if not self.df_musteriler.empty else 1
                yeni_musteri = {'musteri_id': musteri_id, 'isim': musteri_isim, 'soyisim': musteri_soyisim, 'telefon_no': telefon_no, 'adres': ''}
                yeni_satir = pd.DataFrame([yeni_musteri])
                self.df_musteriler = pd.concat([self.df_musteriler, yeni_satir], ignore_index=True)
                self.db_kaydet(self.df_musteriler, self.musteriler_db_path)

        toplam_kar = sum((item['satis_fiyati'] - item['alis_fiyati']) * item['miktar'] for item in self.sepet.values())
        satici = self.satici_menu.get()
        
        if self.current_satis_id:
            self.df_satislar.loc[self.df_satislar['satis_id'] == self.current_satis_id, 'tarih'] = datetime.now()
            self.df_satislar.loc[self.df_satislar['satis_id'] == self.current_satis_id, 'toplam_tutar'] = toplam_tutar
            self.df_satislar.loc[self.df_satislar['satis_id'] == self.current_satis_id, 'indirim_miktari_tl'] = indirim_miktari
            self.df_satislar.loc[self.df_satislar['satis_id'] == self.current_satis_id, 'toplam_kar'] = toplam_kar
            self.df_satislar.loc[self.df_satislar['satis_id'] == self.current_satis_id, 'urunler'] = json.dumps(self.sepet)
            self.df_satislar.loc[self.df_satislar['satis_id'] == self.current_satis_id, 'odeme_yontemi'] = odeme_yontemi
            self.df_satislar.loc[self.df_satislar['satis_id'] == self.current_satis_id, 'satıcı'] = satici
            self.df_satislar.loc[self.df_satislar['satis_id'] == self.current_satis_id, 'musteri_id'] = musteri_id
            messagebox.showinfo("Başarılı", "Satış başarıyla güncellendi.")
        else:
            satis_id = self.df_satislar['satis_id'].max() + 1 if not self.df_satislar.empty else 1
            satis_kaydi = {'satis_id': satis_id, 'tarih': datetime.now(), 'toplam_tutar': toplam_tutar, 'indirim_miktari_tl': indirim_miktari, 'toplam_kar': toplam_kar, 'urunler': json.dumps(self.sepet), 'odeme_yontemi': odeme_yontemi, 'satıcı': satici, 'musteri_id': musteri_id}
            yeni_satir = pd.DataFrame([satis_kaydi])
            self.df_satislar = pd.concat([self.df_satislar, yeni_satir], ignore_index=True)
            messagebox.showinfo("Başarılı", "Satış başarıyla tamamlandı.")
        
        self.db_kaydet(self.df_satislar, self.satislar_db_path)

        if odeme_yontemi == "Borç":
            borc_kaydi = {'satis_id': satis_id, 'musteri_id': musteri_id, 'tarih': datetime.now(), 'borc_miktari': odenecek_tutar, 'odenmis_miktar': 0, 'taksit_miktari': None, 'taksit_gunu': None, 'urunler': json.dumps(self.sepet)}
            borc_satir = pd.DataFrame([borc_kaydi])
            self.df_borclar = pd.concat([self.df_borclar, borc_satir], ignore_index=True)
            self.db_kaydet(self.df_borclar, self.borclar_db_path)
        elif odeme_yontemi == "Taksitli":
            taksit_miktari_str = simpledialog.askstring("Taksit Bilgisi", "Lütfen taksit miktarını girin:")
            taksit_gunu_str = simpledialog.askstring("Taksit Bilgisi", "Lütfen taksit gününü girin (aylık gün, pl. 5, 10):")
            try:
                taksit_miktari = float(taksit_miktari_str)
                taksit_gunu = int(taksit_gunu_str) if taksit_gunu_str else datetime.now().day
                if taksit_miktari <= 0: raise ValueError
                borc_kaydi = {'satis_id': satis_id, 'musteri_id': musteri_id, 'tarih': datetime.now(), 'borc_miktari': odenecek_tutar, 'odenmis_miktar': 0, 'taksit_miktari': taksit_miktari, 'taksit_gunu': taksit_gunu, 'urunler': json.dumps(self.sepet)}
                borc_satir = pd.DataFrame([borc_kaydi])
                self.df_borclar = pd.concat([self.df_borclar, borc_satir], ignore_index=True)
                self.db_kaydet(self.df_borclar, self.borclar_db_path)
            except (ValueError, TypeError):
                messagebox.showerror("Hata", "Lütfen geçerli taksit bilgileri girin. Satış borçlu olarak kaydedildi.")
                borc_kaydi = {'satis_id': satis_id, 'musteri_id': musteri_id, 'tarih': datetime.now(), 'borc_miktari': odenecek_tutar, 'odenmis_miktar': 0, 'taksit_miktari': None, 'taksit_gunu': None, 'urunler': json.dumps(self.sepet)}
                borc_satir = pd.DataFrame([borc_kaydi])
                self.df_borclar = pd.concat([self.df_borclar, borc_satir], ignore_index=True)
                self.db_kaydet(self.df_borclar, self.borclar_db_path)
        
        self.sepet = {}
        self.current_satis_id = None
        self.create_satis_gecmisi_ekrani()
        
    def satis_sil(self, satis_id):
        cevap = messagebox.askyesno("Onay", "Bu satış kaydı kalıcı olarak silinecek. Stoklar iade edilecek. Emin misiniz?", icon='warning')
        if not cevap: return
        
        satis_kaydi = self.df_satislar[self.df_satislar['satis_id'] == satis_id].iloc[0]
        urunler_dict = json.loads(satis_kaydi['urunler'])

        for barkod, item in urunler_dict.items():
            if barkod in self.df_urunler['barkod'].values:
                self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'] += item['miktar']
        
        self.df_satislar = self.df_satislar[self.df_satislar['satis_id'] != satis_id]
        self.df_borclar = self.df_borclar[self.df_borclar['satis_id'] != satis_id]
        
        self.db_kaydet(self.df_urunler, self.urunler_db_path)
        self.db_kaydet(self.df_satislar, self.satislar_db_path)
        self.db_kaydet(self.df_borclar, self.borclar_db_path)
        
        messagebox.showinfo("Başarılı", "Satış kaydı başarıyla silindi ve stoklar iade edildi.")
        self.create_satis_gecmisi_ekrani()

    def create_satis_gecmisi_ekrani(self):
        self.clear_sag_frame()
        self.history.append(self.create_satis_gecmisi_ekrani)
        ctk.CTkLabel(self.sag_frame, text="Satış Geçmişi", font=("Arial", 24, "bold")).pack(pady=10)
        
        filtre_frame = ctk.CTkFrame(self.sag_frame)
        filtre_frame.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(filtre_frame, text="Ara:", font=("Arial", 14)).pack(side="left", padx=5)
        self.search_satis_entry = ctk.CTkEntry(filtre_frame, placeholder_text="Tarih, ürün veya müşteri...")
        self.search_satis_entry.pack(side="left", padx=5, expand=True, fill="x")
        self.search_satis_entry.bind("<KeyRelease>", self.satis_listesi_goster)

        ctk.CTkLabel(filtre_frame, text="Tarih Aralığı:", font=("Arial", 14)).pack(side="left", padx=5)
        self.start_date_entry = ctk.CTkEntry(filtre_frame, placeholder_text="GG-AA-YYYY")
        self.start_date_entry.pack(side="left", padx=5)
        self.end_date_entry = ctk.CTkEntry(filtre_frame, placeholder_text="GG-AA-YYYY")
        self.end_date_entry.pack(side="left", padx=5)
        ctk.CTkButton(filtre_frame, text="Filtrele", command=self.satis_listesi_goster).pack(side="left", padx=10)
        
        self.satis_list_frame = ctk.CTkScrollableFrame(self.sag_frame, height=500)
        self.satis_list_frame.pack(pady=20, padx=20, fill="both", expand=True)
        self.satis_listesi_goster()
        
    def satis_listesi_goster(self, event=None):
        for widget in self.satis_list_frame.winfo_children():
            widget.destroy()
        
        filtreli_df = self.df_satislar.copy()
        filtreli_df = filtreli_df.sort_values(by='tarih', ascending=False)
        
        arama_metni = self.search_satis_entry.get().lower()
        if arama_metni:
            filtreli_df = filtreli_df[
                filtreli_df['tarih'].astype(str).str.contains(arama_metni, na=False) |
                filtreli_df['urunler'].str.lower().str.contains(arama_metni, na=False) |
                filtreli_df['musteri_id'].isin(self.df_musteriler[self.df_musteriler['isim'].str.lower().str.contains(arama_metni, na=False)]['musteri_id'])
            ]
        
        start_date_str = self.start_date_entry.get()
        end_date_str = self.end_date_entry.get()
        
        if start_date_str:
            try:
                start_date = datetime.strptime(start_date_str, '%d-%m-%Y').date()
                filtreli_df = filtreli_df[filtreli_df['tarih'].dt.date >= start_date]
            except ValueError:
                messagebox.showerror("Hata", "Geçerli bir başlangıç tarihi girin (GG-AA-YYYY).")
                return
        
        if end_date_str:
            try:
                end_date = datetime.strptime(end_date_str, '%d-%m-%Y').date()
                filtreli_df = filtreli_df[filtreli_df['tarih'].dt.date <= end_date]
            except ValueError:
                messagebox.showerror("Hata", "Geçerli bir bitiş tarihi girin (GG-AA-YYYY).")
                return

        if filtreli_df.empty:
            ctk.CTkLabel(self.satis_list_frame, text="Satış kaydı bulunamadı.").pack(pady=10)
            return
        
        for _, row in filtreli_df.iterrows():
            satis_frame = ctk.CTkFrame(self.satis_list_frame)
            satis_frame.pack(pady=5, fill="x", padx=10)
            
            tarih_formatli = row['tarih'].strftime("%d-%m-%Y %H:%M")
            musteri_adi = self.df_musteriler[self.df_musteriler['musteri_id'] == row['musteri_id']].iloc[0]['isim'] if not pd.isna(row['musteri_id']) and not self.df_musteriler[self.df_musteriler['musteri_id'] == row['musteri_id']].empty else "Müşteri Yok"
            
            ctk.CTkLabel(satis_frame, text=f"Satış ID: {row['satis_id']} | Tarih: {tarih_formatli} | Tutar: {row['toplam_tutar']:.2f} TL | Ödeme: {row['odeme_yontemi']} | Satıcı: {row['satıcı']} | Müşteri: {musteri_adi}", font=("Arial", 14)).pack(side="left", padx=10)
            
            sil_btn = ctk.CTkButton(satis_frame, text="Sil", width=60, fg_color="#e74c3c", command=lambda r=row: self.satis_sil(r['satis_id']))
            sil_btn.pack(side="right", padx=5)
            
            duzenle_btn = ctk.CTkButton(satis_frame, text="Düzenle", width=60, command=lambda r=row: self.create_satis_form(satis_data=r))
            duzenle_btn.pack(side="right", padx=5)

    def create_borc_yonetim_ekrani(self):
        self.clear_sag_frame()
        self.history.append(self.create_borc_yonetim_ekrani)
        
        top_frame = ctk.CTkFrame(self.sag_frame, fg_color="transparent")
        top_frame.pack(pady=10, fill="x")
        ctk.CTkLabel(top_frame, text="Borç ve Taksit Yönetimi", font=("Arial", 24, "bold")).pack(side="left", padx=20)
        ctk.CTkButton(top_frame, text="Yeni Borç Ekle", command=self.borc_ekle_form).pack(side="right", padx=20)
        
        search_frame = ctk.CTkFrame(self.sag_frame, fg_color="transparent")
        search_frame.pack(pady=5, fill="x", padx=20)
        ctk.CTkLabel(search_frame, text="Müşteri Ara:", font=("Arial", 14)).pack(side="left", padx=5)
        self.musteri_borc_search_entry = ctk.CTkEntry(search_frame, placeholder_text="Müşteri adı veya soyadı...")
        self.musteri_borc_search_entry.pack(side="left", padx=5, expand=True, fill="x")
        self.musteri_borc_search_entry.bind("<KeyRelease>", self.borc_listesi_goster)

        self.borclar_list_frame = ctk.CTkScrollableFrame(self.sag_frame, height=500)
        self.borclar_list_frame.pack(pady=20, padx=20, fill="both", expand=True)
        
        self.borc_listesi_goster()

    def borc_ekle_form(self):
        self.clear_sag_frame()
        self.history.append(self.borc_ekle_form)
        ctk.CTkLabel(self.sag_frame, text="Yeni Borç Ekle", font=("Arial", 24, "bold")).pack(pady=10)
        
        form_frame = ctk.CTkFrame(self.sag_frame)
        form_frame.pack(pady=20, padx=20, fill="x")
        
        ctk.CTkLabel(form_frame, text="Müşteri Adı Soyadı:", font=("Arial", 14)).grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.borc_musteri_isim_entry = ctk.CTkEntry(form_frame)
        self.borc_musteri_isim_entry.grid(row=0, column=1, padx=5, pady=5)
        self.borc_musteri_soyisim_entry = ctk.CTkEntry(form_frame)
        self.borc_musteri_soyisim_entry.grid(row=0, column=2, padx=5, pady=5)

        ctk.CTkLabel(form_frame, text="Telefon No:", font=("Arial", 14)).grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.borc_telefon_entry = ctk.CTkEntry(form_frame)
        self.borc_telefon_entry.grid(row=1, column=1, columnspan=2, padx=5, pady=5, sticky="ew")

        ctk.CTkLabel(form_frame, text="Ödeme Yöntemi:", font=("Arial", 14)).grid(row=2, column=0, padx=5, pady=5, sticky="w")
        self.borc_odeme_menu = ctk.CTkOptionMenu(form_frame, values=["Borç", "Taksitli"])
        self.borc_odeme_menu.grid(row=2, column=1, columnspan=2, padx=5, pady=5, sticky="ew")
        
        self.borc_sepet = {}
        
        urun_arama_frame = ctk.CTkFrame(self.sag_frame, fg_color="transparent")
        urun_arama_frame.pack(pady=10, padx=20, fill="x")
        ctk.CTkLabel(urun_arama_frame, text="Ürün Ekle:", font=("Arial", 14)).pack(side="left", padx=5)
        self.borc_urun_arama_entry = ctk.CTkEntry(urun_arama_frame, placeholder_text="Ürün adı veya barkod...")
        self.borc_urun_arama_entry.pack(side="left", padx=5, expand=True, fill="x")
        self.borc_urun_arama_entry.bind("<KeyRelease>", self.urun_listesi_borc)
        
        self.urun_listesi_borc_frame = ctk.CTkScrollableFrame(self.sag_frame, height=200)
        self.urun_listesi_borc_frame.pack(pady=5, padx=20, fill="x")
        
        self.borc_sepet_frame = ctk.CTkScrollableFrame(self.sag_frame, height=200)
        self.borc_sepet_frame.pack(pady=10, padx=20, fill="x")
        self.borc_sepet_goster()
        
        ctk.CTkButton(self.sag_frame, text="Borcu Kaydet", command=self.borc_kaydet).pack(pady=10)

    def urun_listesi_borc(self, event=None):
        for widget in self.urun_listesi_borc_frame.winfo_children():
            widget.destroy()
        
        arama_metni = self.borc_urun_arama_entry.get().lower()
        
        filtreli_df = self.df_urunler[
            self.df_urunler['urun_adi'].str.lower().str.contains(arama_metni, na=False) |
            self.df_urunler['barkod'].str.lower().str.contains(arama_metni, na=False)
        ]
        
        for _, row in filtreli_df.iterrows():
            urun_frame = ctk.CTkFrame(self.urun_listesi_borc_frame)
            urun_frame.pack(pady=2, fill="x", padx=5)
            ctk.CTkLabel(urun_frame, text=f"{row['urun_adi']} ({row['barkod']}) - {row['satis_fiyati']:.2f} TL (Stok: {int(row['stok_miktari'])})").pack(side="left", expand=True)
            ctk.CTkButton(urun_frame, text="Ekle", command=lambda r=row: self.urun_getir_borc(r)).pack(side="right")
    
    def urun_getir_borc(self, urun):
        barkod = urun['barkod']
        
        if barkod in self.borc_sepet:
            self.borc_sepet[barkod]['miktar'] += 1
        else:
            self.borc_sepet[barkod] = {'urun_adi': urun['urun_adi'], 'satis_fiyati': urun['satis_fiyati'], 'miktar': 1, 'alis_fiyati': urun['alis_fiyati']}

        self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'] -= 1
        self.db_kaydet(self.df_urunler, self.urunler_db_path)
        
        log_kaydi = {'tarih': datetime.now(), 'barkod': barkod, 'miktar_degisimi': -1, 'aciklama': "Borç satışı (Ürün Ekleme)"}
        log_satir = pd.DataFrame([log_kaydi])
        self.df_stok_log = pd.concat([self.df_stok_log, log_satir], ignore_index=True)
        self.db_kaydet(self.df_stok_log, self.stok_log_db_path)

        self.borc_sepet_goster()
        self.borc_urun_arama_entry.delete(0, 'end')
        self.urun_listesi_borc(None)
    
    def borc_sepet_goster(self):
        for widget in self.borc_sepet_frame.winfo_children():
            widget.destroy()
        if not self.borc_sepet:
            ctk.CTkLabel(self.borc_sepet_frame, text="Henüz ürün eklenmedi.").pack()
            return
        
        toplam_borc = 0
        for barkod, item in self.borc_sepet.items():
            urun_frame = ctk.CTkFrame(self.borc_sepet_frame)
            urun_frame.pack(pady=2, fill="x", padx=10)
            
            ctk.CTkLabel(urun_frame, text=item['urun_adi'], font=("Arial", 14)).pack(side="left", padx=5, expand=True)
            
            ctk.CTkLabel(urun_frame, text="Adet:", font=("Arial", 14)).pack(side="left", padx=5)
            miktar_entry = ctk.CTkEntry(urun_frame, width=50)
            miktar_entry.insert(0, str(item['miktar']))
            miktar_entry.pack(side="left", padx=5)
            miktar_entry.bind("<Return>", lambda e, b=barkod: self.borc_sepet_miktar_revize(b, miktar_entry.get()))
            
            ctk.CTkLabel(urun_frame, text="Fiyat:", font=("Arial", 14)).pack(side="left", padx=5)
            fiyat_entry = ctk.CTkEntry(urun_frame, width=80)
            fiyat_entry.insert(0, f"{item['satis_fiyati']:.2f}")
            fiyat_entry.pack(side="left", padx=5)
            fiyat_entry.bind("<KeyRelease>", lambda e, b=barkod: self.borc_sepet_fiyat_revize(b, fiyat_entry.get()))

            sil_btn = ctk.CTkButton(urun_frame, text="Sil", width=50, command=lambda b=barkod: self.borc_sepetten_sil(b))
            sil_btn.pack(side="right", padx=5)
            
            toplam_borc += item['satis_fiyati'] * item['miktar']
        
        ctk.CTkLabel(self.borc_sepet_frame, text=f"Toplam Borç: {toplam_borc:.2f} TL", font=("Arial", 16, "bold")).pack(pady=5)
        
    def borc_sepet_miktar_revize(self, barkod, yeni_miktar_str):
        try:
            yeni_miktar = int(yeni_miktar_str)
            if yeni_miktar < 0: raise ValueError
            eski_miktar = self.borc_sepet[barkod]['miktar']
            stok_degisimi = yeni_miktar - eski_miktar
            urun_stok = self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'].iloc[0]
            if urun_stok - stok_degisimi < 0:
                messagebox.showwarning("Uyarı", "Stok yeterli değil.")
                self.borc_sepet_goster()
                return
            self.borc_sepet[barkod]['miktar'] = yeni_miktar
            self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'] -= stok_degisimi
            self.db_kaydet(self.df_urunler, self.urunler_db_path)
            self.borc_sepet_goster()
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli bir sayı girin.")
            self.borc_sepet_goster()

    def borc_sepet_fiyat_revize(self, barkod, yeni_fiyat_str):
        try:
            yeni_fiyat = float(yeni_fiyat_str)
            if yeni_fiyat < 0: raise ValueError
            self.borc_sepet[barkod]['satis_fiyati'] = yeni_fiyat
            self.borc_sepet_goster()
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli bir fiyat girin.")
            self.borc_sepet_goster()

    def borc_sepetten_sil(self, barkod):
        if barkod in self.borc_sepet:
            miktar_geri = self.borc_sepet[barkod]['miktar']
            self.df_urunler.loc[self.df_urunler['barkod'] == barkod, 'stok_miktari'] += miktar_geri
            del self.borc_sepet[barkod]
            self.db_kaydet(self.df_urunler, self.urunler_db_path)
            self.borc_sepet_goster()

    def borc_kaydet(self):
        try:
            musteri_isim = self.borc_musteri_isim_entry.get().strip()
            musteri_soyisim = self.borc_musteri_soyisim_entry.get().strip()
            telefon_no = self.borc_telefon_entry.get().strip()
            odeme_yontemi = self.borc_odeme_menu.get()

            if not musteri_isim or not musteri_soyisim:
                messagebox.showwarning("Uyarı", "Müşteri adı ve soyadı zorunludur.")
                return
            if not self.borc_sepet:
                messagebox.showwarning("Uyarı", "Lütfen en az bir ürün ekleyin.")
                return
            
            musteri_id = None
            filtre = self.df_musteriler[(self.df_musteriler['isim'] == musteri_isim) & (self.df_musteriler['soyisim'] == musteri_soyisim)]
            if not filtre.empty:
                musteri_id = filtre.iloc[0]['musteri_id']
            else:
                musteri_id = self.df_musteriler['musteri_id'].max() + 1 if not self.df_musteriler.empty else 1
                yeni_musteri = {'musteri_id': musteri_id, 'isim': musteri_isim, 'soyisim': musteri_soyisim, 'telefon_no': telefon_no, 'adres': ''}
                yeni_satir = pd.DataFrame([yeni_musteri])
                self.df_musteriler = pd.concat([self.df_musteriler, yeni_satir], ignore_index=True)
                self.db_kaydet(self.df_musteriler, self.musteriler_db_path)
            
            toplam_borc = sum(item['satis_fiyati'] * item['miktar'] for item in self.borc_sepet.values())
            
            satis_id = self.df_borclar['satis_id'].max() + 1 if not self.df_borclar.empty else 1
            
            taksit_miktari = None
            taksit_gunu = None
            if odeme_yontemi == "Taksitli":
                taksit_miktari_str = simpledialog.askstring("Taksit Bilgisi", "Lütfen taksit miktarını girin:", parent=self)
                taksit_gunu_str = simpledialog.askstring("Taksit Bilgisi", "Lütfen taksit gününü girin (aylık gün, pl. 5, 10):", parent=self)
                try:
                    taksit_miktari = float(taksit_miktari_str)
                    taksit_gunu = int(taksit_gunu_str) if taksit_gunu_str else None
                except (ValueError, TypeError):
                    messagebox.showerror("Hata", "Lütfen geçerli taksit bilgileri girin. Borç olarak kaydediliyor.")
                    odeme_yontemi = "Borç"

            yeni_borc = {'satis_id': satis_id, 'musteri_id': musteri_id, 'tarih': datetime.now(), 'borc_miktari': toplam_borc, 'odenmis_miktar': 0, 'taksit_miktari': taksit_miktari, 'taksit_gunu': taksit_gunu, 'urunler': json.dumps(self.borc_sepet)}
            yeni_satir = pd.DataFrame([yeni_borc])
            self.df_borclar = pd.concat([self.df_borclar, yeni_satir], ignore_index=True)
            self.db_kaydet(self.df_borclar, self.borclar_db_path)
            
            self.borc_sepet = {}
            messagebox.showinfo("Başarılı", "Borç kaydı başarıyla eklendi.")
            self.create_borc_yonetim_ekrani()
        except ValueError:
            messagebox.showerror("Hata", "Lütfen geçerli sayısal değer girin.")

    def borc_listesi_goster(self, event=None):
        for widget in self.borclar_list_frame.winfo_children():
            widget.destroy()
        
        arama_metni = self.musteri_borc_search_entry.get().lower()
        
        if arama_metni:
            filtreli_musteri_ids = self.df_musteriler[
                self.df_musteriler['isim'].str.lower().str.contains(arama_metni, na=False) |
                self.df_musteriler['soyisim'].str.lower().str.contains(arama_metni, na=False)
            ]['musteri_id'].tolist()
            filtreli_df = self.df_borclar[self.df_borclar['musteri_id'].isin(filtreli_musteri_ids)]
        else:
            filtreli_df = self.df_borclar.copy()

        if filtreli_df.empty:
            ctk.CTkLabel(self.borclar_list_frame, text="Henüz borç kaydı bulunmamaktadır.").pack(pady=10)
            return
        
        bugunun_gunu = datetime.now().day
        
        for index, row in filtreli_df.iterrows():
            kalan_borc = row['borc_miktari'] - row['odenmis_miktar']
            tarih_formatli = row['tarih'].strftime("%d-%m-%Y")
            musteri_adi_soyadi = f"{self.df_musteriler[self.df_musteriler['musteri_id'] == row['musteri_id']].iloc[0]['isim']} {self.df_musteriler[self.df_musteriler['musteri_id'] == row['musteri_id']].iloc[0]['soyisim']}"
            
            bg_color = None
            if not pd.isna(row['taksit_gunu']) and row['taksit_gunu'] == bugunun_gunu:
                bg_color = "red"
            
            borc_frame = ctk.CTkFrame(self.borclar_list_frame, fg_color=bg_color)
            borc_frame.pack(pady=5, fill="x", padx=10)
            
            ctk.CTkLabel(borc_frame, text=f"Müşteri: {musteri_adi_soyadi} | Tarih: {tarih_formatli} | Borç: {row['borc_miktari']:.2f} TL | Ödenen: {row['odenmis_miktar']:.2f} TL | Kalan: {kalan_borc:.2f} TL", font=("Arial", 14)).pack(side="left", padx=10)
            
            sil_btn = ctk.CTkButton(borc_frame, text="Sil", width=60, fg_color="#e74c3c", command=lambda r=row: self.borc_sil(r.name))
            sil_btn.pack(side="right", padx=5)
            
            odeme_btn = ctk.CTkButton(borc_frame, text="Ödeme Ekle", width=60, command=lambda r=row: self.borc_odeme_ekle(r.name))
            odeme_btn.pack(side="right", padx=5)

            detay_btn = ctk.CTkButton(borc_frame, text="Detaylar", width=60, command=lambda r=row: self.borc_detaylari(r.name))
            detay_btn.pack(side="right", padx=5)
    
    def borc_detaylari(self, index):
        self.clear_sag_frame()
        self.history.append(self.borc_detaylari)
        row = self.df_borclar.iloc[index]
        musteri_adi_soyadi = f"{self.df_musteriler[self.df_musteriler['musteri_id'] == row['musteri_id']].iloc[0]['isim']} {self.df_musteriler[self.df_musteriler['musteri_id'] == row['musteri_id']].iloc[0]['soyisim']}"
        ctk.CTkLabel(self.sag_frame, text=f"{musteri_adi_soyadi} Borç Detayları", font=("Arial", 24, "bold")).pack(pady=10)
        
        ctk.CTkLabel(self.sag_frame, text=f"Borç Miktarı: {row['borc_miktari']:.2f} TL", font=("Arial", 16)).pack(pady=5)
        ctk.CTkLabel(self.sag_frame, text=f"Ödenen: {row['odenmis_miktar']:.2f} TL", font=("Arial", 16)).pack(pady=5)
        ctk.CTkLabel(self.sag_frame, text=f"Kalan: {row['borc_miktari'] - row['odenmis_miktar']:.2f} TL", font=("Arial", 16)).pack(pady=5)

        urunler = json.loads(row['urunler'])
        if urunler:
            ctk.CTkLabel(self.sag_frame, text="İlgili Ürünler:", font=("Arial", 16, "bold")).pack(pady=10)
            urunler_frame = ctk.CTkScrollableFrame(self.sag_frame, height=200)
            urunler_frame.pack(fill="x", padx=20)
            for barkod, urun_data in urunler.items():
                ctk.CTkLabel(urunler_frame, text=f"- {urun_data['urun_adi']} (Adet: {urun_data['miktar']}, Fiyat: {urun_data['satis_fiyati']:.2f} TL)").pack(fill="x", padx=10)
        
        odemeler = self.df_borc_odemeler[self.df_borc_odemeler['satis_id'] == row['satis_id']]
        ctk.CTkLabel(self.sag_frame, text="Ödeme Kayıtları:", font=("Arial", 16, "bold")).pack(pady=10)
        odemeler_frame = ctk.CTkScrollableFrame(self.sag_frame, height=200)
        odemeler_frame.pack(fill="x", padx=20)
        if odemeler.empty:
            ctk.CTkLabel(odemeler_frame, text="Henüz ödeme yapılmamış.").pack(pady=5)
        else:
            for _, odeme_row in odemeler.iterrows():
                tarih = odeme_row['tarih'].strftime('%d-%m-%Y %H:%M')
                ctk.CTkLabel(odemeler_frame, text=f"- {tarih}: {odeme_row['miktar']:.2f} TL").pack(fill="x", padx=10)
        
        ctk.CTkButton(self.sag_frame, text="Geri Dön", command=self.create_borc_yonetim_ekrani).pack(pady=10)
    
    def borc_sil(self, index):
        cevap = messagebox.askyesno("Onay", "Bu borç kaydı silinecek. Emin misiniz?", icon='warning')
        if not cevap: return
        satis_id = self.df_borclar.loc[index, 'satis_id']
        self.df_borclar.drop(index, inplace=True)
        self.df_borclar.reset_index(drop=True, inplace=True)
        self.df_borc_odemeler = self.df_borc_odemeler[self.df_borc_odemeler['satis_id'] != satis_id]
        self.db_kaydet(self.df_borclar, self.borclar_db_path)
        self.db_kaydet(self.df_borc_odemeler, self.borc_odemeler_db_path)
        messagebox.showinfo("Başarılı", "Borç kaydı silindi.")
        self.create_borc_yonetim_ekrani()

    def borc_odeme_ekle(self, index):
        odeme_miktari_str = simpledialog.askstring("Ödeme Ekle", "Lütfen ödeme miktarını girin:", parent=self)
        if odeme_miktari_str:
            try:
                odeme_miktari = float(odeme_miktari_str)
                row = self.df_borclar.iloc[index]
                kalan_borc = row['borc_miktari'] - row['odenmis_miktar']
                if odeme_miktari > kalan_borc:
                    messagebox.showwarning("Uyarı", "Ödeme miktarı kalan borçtan fazla olamaz.")
                    return
                
                self.df_borclar.loc[index, 'odenmis_miktar'] += odeme_miktari
                self.db_kaydet(self.df_borclar, self.borclar_db_path)
                
                odeme_id = self.df_borc_odemeler['odeme_id'].max() + 1 if not self.df_borc_odemeler.empty else 1
                odeme_kaydi = {'odeme_id': odeme_id, 'satis_id': row['satis_id'], 'tarih': datetime.now(), 'miktar': odeme_miktari, 'aciklama': "Ödeme yapıldı"}
                odeme_satir = pd.DataFrame([odeme_kaydi])
                self.df_borc_odemeler = pd.concat([self.df_borc_odemeler, odeme_satir], ignore_index=True)
                self.db_kaydet(self.df_borc_odemeler, self.borc_odemeler_db_path)
                
                messagebox.showinfo("Başarılı", "Ödeme başarıyla eklendi.")
                self.create_borc_yonetim_ekrani()
            except ValueError:
                messagebox.showerror("Hata", "Geçerli bir sayı girin.")
    
    def create_yonetici_ekrani_password(self):
        self.clear_sag_frame()
        self.history.append(self.create_yonetici_ekrani_password)
        ctk.CTkLabel(self.sag_frame, text="Yönetici Paneli Giriş", font=("Arial", 24, "bold")).pack(pady=20)
        
        sifre_frame = ctk.CTkFrame(self.sag_frame)
        sifre_frame.pack(pady=10)
        ctk.CTkLabel(sifre_frame, text="Şifre:", font=("Arial", 16)).pack(side="left", padx=10)
        self.sifre_entry = ctk.CTkEntry(sifre_frame, show='*')
        self.sifre_entry.pack(side="left", padx=10)
        
        ctk.CTkButton(self.sag_frame, text="Giriş Yap", command=self.create_yonetici_ekrani).pack(pady=10)

    def create_yonetici_ekrani(self):
        if self.sifre_entry.get() == "1234":
            self.clear_sag_frame()
            self.history.append(self.create_yonetici_ekrani)
            ctk.CTkLabel(self.sag_frame, text="Yönetici Paneli", font=("Arial", 24, "bold")).pack(pady=10)
            ctk.CTkButton(self.sag_frame, text="Ciro ve Kar Raporu", command=self.create_ciro_raporu).pack(pady=10)
            ctk.CTkButton(self.sag_frame, text="Yıllık Devir Sistemi", command=self.devir_sistemi_password).pack(pady=10)
            ctk.CTkButton(self.sag_frame, text="Veritabanı Yönetimi", command=self.create_db_yonetimi_menu).pack(pady=10)
        else:
            messagebox.showerror("Hata", "Yanlış şifre.")

    def devir_sistemi_password(self):
        sifre = simpledialog.askstring("Yıllık Devir", "Devir işlemini başlatmak için şifreyi (0000) girin:", show='*')
        if sifre == "0000":
            self.devir_sistemi()
        else:
            messagebox.showerror("Hata", "Yanlış şifre.")

    def devir_sistemi(self):
        cevap = messagebox.askyesno("Yıllık Devir", "Yıllık devir işlemini başlatmak ister misiniz? Tüm satış ve borç kayıtları sıfırlanacaktır. Önce bir yedek oluşturulacaktır.")
        if not cevap: return

        try:
            document = Document()
            document.add_heading('Yıllık Devir Raporu', 0)
            
            document.add_heading('Satış Geçmişi', level=1)
            satis_table = document.add_table(self.df_satislar.shape[0]+1, self.df_satislar.shape[1])
            for j in range(self.df_satislar.shape[-1]):
                satis_table.cell(0, j).text = self.df_satislar.columns[j]
            for i in range(self.df_satislar.shape[0]):
                for j in range(self.df_satislar.shape[-1]):
                    satis_table.cell(i+1, j).text = str(self.df_satislar.values[i,j])
            
            document.add_heading('Borç Kayıtları', level=1)
            borc_table = document.add_table(self.df_borclar.shape[0]+1, self.df_borclar.shape[1])
            for j in range(self.df_borclar.shape[-1]):
                borc_table.cell(0, j).text = self.df_borclar.columns[j]
            for i in range(self.df_borclar.shape[0]):
                for j in range(self.df_borclar.shape[-1]):
                    borc_table.cell(i+1, j).text = str(self.df_borclar.values[i,j])

            filename = f"Devir_Yedek_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            document.save(filename)
            messagebox.showinfo("Başarılı", f"Devir raporu '{filename}' dosyasına kaydedildi.")

            self.df_satislar = pd.DataFrame(columns=self.df_satislar.columns)
            self.df_borclar = pd.DataFrame(columns=self.df_borclar.columns)
            self.db_kaydet(self.df_satislar, self.satislar_db_path)
            self.db_kaydet(self.df_borclar, self.borclar_db_path)
            messagebox.showinfo("Başarılı", "Yıllık devir işlemi tamamlandı. Veritabanları sıfırlandı.")
            self.create_main_menu()
        except Exception as e:
            messagebox.showerror("Hata", f"Devir işlemi sırasında bir hata oluştu: {e}")

    def create_ciro_raporu(self):
        self.clear_sag_frame()
        self.history.append(self.create_ciro_raporu)
        ctk.CTkLabel(self.sag_frame, text="Satış ve Kar Raporu", font=("Arial", 20, "bold")).pack(pady=10)
        
        filtre_frame = ctk.CTkFrame(self.sag_frame)
        filtre_frame.pack(pady=10, padx=20, fill="x")
        
        ctk.CTkButton(filtre_frame, text="Günlük", command=lambda: self.goster_rapor('gunluk')).pack(side="left", padx=5)
        ctk.CTkButton(filtre_frame, text="Aylık", command=lambda: self.goster_rapor('aylik')).pack(side="left", padx=5)
        ctk.CTkButton(filtre_frame, text="Yıllık", command=lambda: self.goster_rapor('yillik')).pack(side="left", padx=5)
        
        ctk.CTkLabel(filtre_frame, text="Başlangıç Tarihi:").pack(side="left", padx=5)
        self.start_date_entry = ctk.CTkEntry(filtre_frame, placeholder_text="GG-AA-YYYY")
        self.start_date_entry.pack(side="left", padx=5)
        ctk.CTkLabel(filtre_frame, text="Bitiş Tarihi:").pack(side="left", padx=5)
        self.end_date_entry = ctk.CTkEntry(filtre_frame, placeholder_text="GG-AA-YYYY")
        self.end_date_entry.pack(side="left", padx=5)
        ctk.CTkButton(filtre_frame, text="Özel Filtrele", command=lambda: self.goster_rapor('ozel')).pack(side="left", padx=10)

        self.rapor_list_frame = ctk.CTkScrollableFrame(self.sag_frame, height=300)
        self.rapor_list_frame.pack(fill="x", expand=False, padx=20)
        
        self.rapor_graph_frame = ctk.CTkFrame(self.sag_frame)
        self.rapor_graph_frame.pack(pady=10, padx=20, fill="both", expand=True)

        self.goster_rapor('gunluk')

    def goster_rapor(self, filtre_tipi='gunluk'):
        for widget in self.rapor_list_frame.winfo_children():
            widget.destroy()
        for widget in self.rapor_graph_frame.winfo_children():
            widget.destroy()

        filtered_df = self.df_satislar.copy()
        bugun = datetime.now()

        if filtre_tipi == 'gunluk':
            start_date = bugun.date()
            end_date = bugun.date()
        elif filtre_tipi == 'aylik':
            start_date = bugun.replace(day=1).date()
            end_date = bugun.date()
        elif filtre_tipi == 'yillik':
            start_date = bugun.replace(month=1, day=1).date()
            end_date = bugun.date()
        elif filtre_tipi == 'ozel':
            try:
                start_date = datetime.strptime(self.start_date_entry.get(), '%d-%m-%Y').date()
                end_date = datetime.strptime(self.end_date_entry.get(), '%d-%m-%Y').date()
            except ValueError:
                messagebox.showerror("Hata", "Geçerli bir başlangıç ve bitiş tarihi girin (GG-AA-YYYY).")
                return
        
        filtered_df = filtered_df[filtered_df['tarih'].dt.date >= start_date]
        filtered_df = filtered_df[filtered_df['tarih'].dt.date <= end_date]

        if filtered_df.empty:
            ctk.CTkLabel(self.rapor_list_frame, text="Belirtilen tarihlerde satış yapılmamış.").pack(pady=20)
            return

        grouped = filtered_df.groupby(filtered_df['tarih'].dt.date).agg(
            toplam_ciro=('toplam_tutar', 'sum'),
            toplam_kar=('toplam_kar', 'sum')
        ).reset_index()
        
        for _, row in grouped.iterrows():
            rapor_frame = ctk.CTkFrame(self.rapor_list_frame)
            rapor_frame.pack(pady=5, fill="x", padx=10)
            ctk.CTkLabel(rapor_frame, text=f"Tarih: {row['tarih']} | Ciro: {row['toplam_ciro']:.2f} TL | Kar: {row['toplam_kar']:.2f} TL", font=("Arial", 14)).pack(pady=5, fill="x", padx=10)

        # Grafik
        fig, ax = plt.subplots(figsize=(8, 4), facecolor='#ebebeb')
        fig.set_facecolor('#ebebeb')
        ax.set_facecolor('#ebebeb')
        ax.tick_params(colors='black')
        ax.spines['bottom'].set_color('black')
        ax.spines['left'].set_color('black')
        ax.set_title("Ciro ve Kar Grafiği", color='black')
        ax.set_xlabel("Tarih", color='black')
        ax.set_ylabel("Miktar (TL)", color='black')
        
        ax.plot(grouped['tarih'], grouped['toplam_ciro'], label="Ciro", marker='o')
        ax.plot(grouped['tarih'], grouped['toplam_kar'], label="Kar", marker='o')
        ax.legend()
        plt.xticks(rotation=45)
        plt.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=self.rapor_graph_frame)
        canvas_widget = canvas.get_tk_widget()
        canvas_widget.pack(fill=tk.BOTH, expand=True)


    def create_stok_raporu(self):
        self.clear_sag_frame()
        self.history.append(self.create_stok_raporu)
        ctk.CTkLabel(self.sag_frame, text="Stok Durum Raporu", font=("Arial", 20, "bold")).pack(pady=10)
        
        tabview = ctk.CTkTabview(self.sag_frame)
        tabview.pack(pady=10, padx=20, fill="both", expand=True)
        tabview.add("Mevcut Stok")
        tabview.add("Stok Değişim Geçmişi")

        self.goster_mevcut_stok(tabview.tab("Mevcut Stok"))
        self.goster_stok_gecmisi(tabview.tab("Stok Değişim Geçmişi"))

    def goster_mevcut_stok(self, parent_frame):
        stok_rapor_frame = ctk.CTkScrollableFrame(parent_frame, height=600)
        stok_rapor_frame.pack(fill="both", expand=True, padx=20)
        if self.df_urunler.empty:
            ctk.CTkLabel(stok_rapor_frame, text="Stokta ürün bulunamadı.").pack(pady=20)
            return
        self.df_urunler = self.df_urunler.sort_values(by='stok_miktari', ascending=True)
        for _, row in self.df_urunler.iterrows():
            renk = "red" if row['stok_miktari'] < 5 else "green"
            urun_frame = ctk.CTkFrame(stok_rapor_frame)
            urun_frame.pack(pady=5, fill="x", padx=10)
            ctk.CTkLabel(urun_frame, text=f"{row['urun_adi']} ({row['barkod']})", font=("Arial", 16, "bold")).pack(side="left", padx=10)
            ctk.CTkLabel(urun_frame, text=f"Stok: {int(row['stok_miktari'])}", font=("Arial", 14), text_color=renk).pack(side="right", padx=10)
            ctk.CTkLabel(urun_frame, text=f"Satış Fiyatı: {row['satis_fiyati']:.2f} TL", font=("Arial", 14)).pack(side="right", padx=10)
            ctk.CTkLabel(urun_frame, text=f"Alış Fiyatı: {row['alis_fiyati']:.2f} TL", font=("Arial", 14)).pack(side="right", padx=10)

    def goster_stok_gecmisi(self, parent_frame):
        stok_log_frame = ctk.CTkScrollableFrame(parent_frame, height=600)
        stok_log_frame.pack(fill="both", expand=True, padx=20)
        if self.df_stok_log.empty:
            ctk.CTkLabel(stok_log_frame, text="Stok değişim kaydı bulunamadı.").pack(pady=20)
            return
        self.df_stok_log = self.df_stok_log.sort_values(by='tarih', ascending=False)
        for _, row in self.df_stok_log.iterrows():
            urun_adi = self.df_urunler[self.df_urunler['barkod'] == row['barkod']]['urun_adi'].iloc[0] if row['barkod'] in self.df_urunler['barkod'].values else "Bilinmeyen Ürün"
            ctk.CTkLabel(stok_log_frame, text=f"Tarih: {row['tarih'].strftime('%d-%m-%Y %H:%M')} | Ürün: {urun_adi} | Barkod: {row['barkod']} | Miktar: {row['miktar_degisimi']} | Açıklama: {row['aciklama']}").pack(pady=2, fill="x", padx=10)

    def barkod_olustur_ve_kaydet(self, barkod, urun_adi, fiyat):
        try:
            os.makedirs("barkodlar", exist_ok=True)
            filename = f"barkodlar/{barkod}"
            code128 = barcode.get_barcode_class('code128')
            writer = ImageWriter()
            my_barcode = code128(barkod, writer=writer)
            my_barcode.save(filename)
            img = Image.open(f"{filename}.png")
            genislik, yukseklik = img.size
            
            yazi_yuksekligi = 60
            yeni_yukseklik = yukseklik + yazi_yuksekligi
            yeni_img = Image.new('RGB', (genislik, yeni_yukseklik), 'white')
            yeni_img.paste(img, (0, 0))
            
            draw = ImageDraw.Draw(yeni_img)
            
            try:
                font_yol = "arial.ttf"
                font = ImageFont.truetype(font_yol, 20)
            except IOError:
                font_yol = "C:/Windows/Fonts/arial.ttf"
                try:
                    font = ImageFont.truetype(font_yol, 20)
                except IOError:
                    messagebox.showwarning("Uyarı", "arial.ttf font dosyası bulunamadı. Metin eklenemedi.")
                    yeni_img.save(f"{filename}.png")
                    return

            bbox_urun_adi = draw.textbbox((0,0), urun_adi, font=font)
            text_genislik_urun_adi = bbox_urun_adi[2] - bbox_urun_adi[0]
            x_konum_urun_adi = (genislik - text_genislik_urun_adi) / 2
            draw.text((x_konum_urun_adi, yukseklik + 5), urun_adi, font=font, fill=(0, 0, 0))

            fiyat_text = f"Fiyat: {fiyat:.2f} TL"
            bbox_fiyat = draw.textbbox((0,0), fiyat_text, font=font)
            text_genislik_fiyat = bbox_fiyat[2] - bbox_fiyat[0]
            x_konum_fiyat = (genislik - text_genislik_fiyat) / 2
            draw.text((x_konum_fiyat, yukseklik + 30), fiyat_text, font=font, fill=(0, 0, 0))

            yeni_img.save(f"{filename}.png")
            messagebox.showinfo("Başarılı", "Barkod etiketi 'barkodlar' klasörüne kaydedildi.")
            cevap = messagebox.askyesno("Klasörü Aç", "Barkod etiketlerinin bulunduğu klasörü açmak ister misiniz?")
            if cevap: webbrowser.open(f'file:///{os.path.realpath("barkodlar")}')
        except Exception as e:
            messagebox.showerror("Hata", f"Barkod oluşturma sırasında bir hata oluştu: {e}")

    def create_firma_yonetimi(self):
        self.clear_sag_frame()
        self.history.append(self.create_firma_yonetimi)
        ctk.CTkLabel(self.sag_frame, text="Firma Yönetimi", font=("Arial", 24, "bold")).pack(pady=10)
        
        search_frame = ctk.CTkFrame(self.sag_frame, fg_color="transparent")
        search_frame.pack(pady=5, fill="x", padx=20)
        ctk.CTkLabel(search_frame, text="Firma Ara:", font=("Arial", 14)).pack(side="left", padx=5)
        self.firma_search_entry = ctk.CTkEntry(search_frame, placeholder_text="Firma adı...")
        self.firma_search_entry.pack(side="left", padx=5, expand=True, fill="x")
        self.firma_search_entry.bind("<KeyRelease>", self.firma_listesi_goster)
        
        self.firma_list_frame = ctk.CTkScrollableFrame(self.sag_frame, height=500)
        self.firma_list_frame.pack(pady=20, padx=20, fill="both", expand=True)
        self.firma_listesi_goster()
        
    def firma_listesi_goster(self, event=None):
        for widget in self.firma_list_frame.winfo_children():
            widget.destroy()
        
        arama_metni = self.firma_search_entry.get().lower()
        if self.df_firmalar.empty:
            ctk.CTkLabel(self.firma_list_frame, text="Kayıtlı firma bulunamadı.").pack(pady=10)
            return

        filtreli_df = self.df_firmalar[self.df_firmalar['firma_adi'].str.lower().str.contains(arama_metni, na=False)]
        
        for index, row in filtreli_df.iterrows():
            firma_frame = ctk.CTkFrame(self.firma_list_frame)
            firma_frame.pack(pady=5, fill="x", padx=10)
            ctk.CTkLabel(firma_frame, text=row['firma_adi'], font=("Arial", 16, "bold")).pack(side="left", padx=10, expand=True)
            
            ctk.CTkButton(firma_frame, text="Sil", width=60, fg_color="#e74c3c", command=lambda r=row: self.firma_sil(r['firma_adi'])).pack(side="right", padx=5)
            ctk.CTkButton(firma_frame, text="Düzenle", width=60, command=lambda r=row: self.firma_form(edit_mode=True, firma_data=r)).pack(side="right", padx=5)
            ctk.CTkButton(firma_frame, text="Detaylar", width=60, command=lambda r=row: self.firma_detay(r)).pack(side="right", padx=5)
    
    def firma_form(self, edit_mode=False, firma_data=None):
        self.clear_sag_frame()
        self.history.append(lambda: self.firma_form(edit_mode, firma_data))
        
        baslik = "Firma Düzenle" if edit_mode else "Yeni Firma Ekle"
        ctk.CTkLabel(self.sag_frame, text=baslik, font=("Arial", 24, "bold")).pack(pady=10)
        
        form_frame = ctk.CTkFrame(self.sag_frame)
        form_frame.pack(pady=20, padx=20, fill="x")
        
        self.firma_entry_vars = {}
        etiketler = ["Firma Adı:", "Adres:", "Telefon:", "Vergi Dairesi:", "Yetkili Kişi:", "Notlar:"]
        
        for i, etiket in enumerate(etiketler):
            ctk.CTkLabel(form_frame, text=etiket).grid(row=i, column=0, padx=10, pady=5, sticky="w")
            var = ctk.StringVar()
            self.firma_entry_vars[etiket] = var
            entry = ctk.CTkEntry(form_frame, textvariable=var)
            entry.grid(row=i, column=1, padx=10, pady=5, sticky="ew")
            
            if edit_mode:
                var.set(str(firma_data.get(etiket.replace(":", "").replace(" ", "_").lower(), '')))
                if etiket == "Firma Adı:":
                    entry.configure(state='disabled')
        
        ctk.CTkButton(form_frame, text="Kaydet", command=lambda: self.firma_kaydet(edit_mode, firma_data)).grid(row=len(etiketler), column=0, columnspan=2, pady=10)

    def firma_kaydet(self, edit_mode, firma_data=None):
        firma_adi = self.firma_entry_vars["Firma Adı:"].get().strip()
        adres = self.firma_entry_vars["Adres:"].get()
        telefon = self.firma_entry_vars["Telefon:"].get()
        vergi_dairesi = self.firma_entry_vars["Vergi Dairesi:"].get()
        yetkili_kisi = self.firma_entry_vars["Yetkili Kişi:"].get()
        notlar = self.firma_entry_vars["Notlar:"].get()

        if not firma_adi:
            messagebox.showwarning("Uyarı", "Firma Adı boş bırakılamaz.")
            return

        if edit_mode:
            self.df_firmalar.loc[self.df_firmalar['firma_adi'] == firma_data['firma_adi'], 'adres'] = adres
            self.df_firmalar.loc[self.df_firmalar['firma_adi'] == firma_data['firma_adi'], 'telefon'] = telefon
            self.df_firmalar.loc[self.df_firmalar['firma_adi'] == firma_data['firma_adi'], 'vergi_dairesi'] = vergi_dairesi
            self.df_firmalar.loc[self.df_firmalar['firma_adi'] == firma_data['firma_adi'], 'yetkili_kisi'] = yetkili_kisi
            self.df_firmalar.loc[self.df_firmalar['firma_adi'] == firma_data['firma_adi'], 'notlar'] = notlar
            messagebox.showinfo("Başarılı", "Firma bilgileri başarıyla güncellendi.")
        else:
            if firma_adi in self.df_firmalar['firma_adi'].values:
                messagebox.showwarning("Uyarı", "Bu firma adı zaten mevcut.")
                return
            yeni_firma = pd.DataFrame([{'firma_adi': firma_adi, 'adres': adres, 'telefon': telefon, 'vergi_dairesi': vergi_dairesi, 'yetkili_kisi': yetkili_kisi, 'notlar': notlar}])
            self.df_firmalar = pd.concat([self.df_firmalar, yeni_firma], ignore_index=True)
            messagebox.showinfo("Başarılı", "Yeni firma başarıyla kaydedildi.")

        self.db_kaydet(self.df_firmalar, self.firmalar_db_path)
        self.create_firma_yonetimi()

    def firma_kaydet_veya_guncelle(self, firma_adi):
        if firma_adi not in self.df_firmalar['firma_adi'].values:
            yeni_firma = pd.DataFrame([{'firma_adi': firma_adi, 'adres': '', 'telefon': '', 'vergi_dairesi': '', 'yetkili_kisi': '', 'notlar': ''}])
            self.df_firmalar = pd.concat([self.df_firmalar, yeni_firma], ignore_index=True)
            self.db_kaydet(self.df_firmalar, self.firmalar_db_path)

    def firma_detay(self, firma_data):
        self.clear_sag_frame()
        self.history.append(lambda: self.firma_detay(firma_data))
        ctk.CTkLabel(self.sag_frame, text=f"{firma_data['firma_adi']} Detayları", font=("Arial", 24, "bold")).pack(pady=10)
        
        detay_frame = ctk.CTkFrame(self.sag_frame)
        detay_frame.pack(pady=20, padx=20, fill="x")
        
        etiketler = ["Firma Adı:", "Adres:", "Telefon:", "Vergi Dairesi:", "Yetkili Kişi:", "Notlar:"]
        
        for i, etiket in enumerate(etiketler):
            deger = firma_data.get(etiket.replace(":", "").replace(" ", "_").lower(), 'Yok')
            ctk.CTkLabel(detay_frame, text=f"{etiket} {deger}").pack(pady=5, fill="x", padx=10)
        
        ctk.CTkButton(self.sag_frame, text="Geri Dön", command=self.create_firma_yonetimi).pack(pady=10)
    
    def firma_sil(self, firma_adi):
        cevap = messagebox.askyesno("Onay", f"{firma_adi} firması kalıcı olarak silinecek. Emin misiniz?", icon='warning')
        if cevap:
            self.df_firmalar = self.df_firmalar[self.df_firmalar['firma_adi'] != firma_adi]
            self.db_kaydet(self.df_firmalar, self.firmalar_db_path)
            messagebox.showinfo("Başarılı", "Firma başarıyla silindi.")
            self.create_firma_yonetimi()

    def create_db_yonetimi_menu(self):
        self.clear_sag_frame()
        self.history.append(self.create_db_yonetimi_menu)
        ctk.CTkLabel(self.sag_frame, text="Veritabanı Yönetimi", font=("Arial", 24, "bold")).pack(pady=10)
        ctk.CTkButton(self.sag_frame, text="Yedek Al", command=self.yedek_al).pack(pady=10)
        ctk.CTkButton(self.sag_frame, text="Yedekten Geri Yükle", command=self.yedekten_yukle).pack(pady=10)
        ctk.CTkButton(self.sag_frame, text="Veritabanını Dışa Aktar", command=self.db_disa_aktar).pack(pady=10)
        ctk.CTkButton(self.sag_frame, text="Tüm Veritabanını Sıfırla", command=self.db_sifirla).pack(pady=10)

    def yedek_al(self):
        yedek_klasoru = f"yedek_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        os.makedirs(yedek_klasoru, exist_ok=True)
        try:
            shutil.copy(self.urunler_db_path, yedek_klasoru)
            shutil.copy(self.satislar_db_path, yedek_klasoru)
            shutil.copy(self.stok_log_db_path, yedek_klasoru)
            shutil.copy(self.borclar_db_path, yedek_klasoru)
            shutil.copy(self.borc_odemeler_db_path, yedek_klasoru)
            shutil.copy(self.musteriler_db_path, yedek_klasoru)
            shutil.copy(self.firmalar_db_path, yedek_klasoru)
            messagebox.showinfo("Başarılı", f"Veritabanı yedeği '{yedek_klasoru}' klasörüne alındı.")
        except FileNotFoundError:
            messagebox.showwarning("Uyarı", "Veritabanı dosyaları bulunamadı. Henüz veri girişi yapılmamış olabilir.")

    def yedekten_yukle(self):
        dosya_yolu = filedialog.askopenfilename(title="Yedek dosyasını seçin", filetypes=[("CSV Dosyaları", "*.csv")])
        if dosya_yolu:
            try:
                dosya_adi = os.path.basename(dosya_yolu)
                if dosya_adi in [self.urunler_db_path, self.satislar_db_path, self.stok_log_db_path, self.borclar_db_path, self.borc_odemeler_db_path, self.musteriler_db_path, self.firmalar_db_path]:
                    shutil.copy(dosya_yolu, os.path.join(os.getcwd(), dosya_adi))
                    if dosya_adi == self.urunler_db_path: self.df_urunler = self.db_yukle(self.urunler_db_path)
                    elif dosya_adi == self.satislar_db_path: self.df_satislar = self.db_yukle(self.satislar_db_path)
                    elif dosya_adi == self.stok_log_db_path: self.df_stok_log = self.db_yukle(self.stok_log_db_path)
                    elif dosya_adi == self.borclar_db_path: self.df_borclar = self.db_yukle(self.borclar_db_path)
                    elif dosya_adi == self.borc_odemeler_db_path: self.df_borc_odemeler = self.db_yukle(self.borc_odemeler_db_path)
                    elif dosya_adi == self.musteriler_db_path: self.df_musteriler = self.db_yukle(self.musteriler_db_path)
                    elif dosya_adi == self.firmalar_db_path: self.df_firmalar = self.db_yukle(self.firmalar_db_path)
                    messagebox.showinfo("Başarılı", f"Veritabanı dosyası '{dosya_adi}' başarıyla geri yüklendi.")
                    self.create_main_menu()
                else:
                    messagebox.showerror("Hata", "Lütfen geçerli bir yedek dosyası seçin.")
            except Exception as e:
                messagebox.showerror("Hata", f"Geri yükleme sırasında bir hata oluştu: {e}")

    def db_sifirla(self):
        cevap = messagebox.askyesno("Onay", "Tüm veritabanı içeriği silinecek. Emin misiniz?", icon='warning')
        if cevap:
            try:
                open(self.urunler_db_path, 'w').close()
                open(self.satislar_db_path, 'w').close()
                open(self.stok_log_db_path, 'w').close()
                open(self.borclar_db_path, 'w').close()
                open(self.borc_odemeler_db_path, 'w').close()
                open(self.musteriler_db_path, 'w').close()
                open(self.firmalar_db_path, 'w').close()
            except: pass
            self.df_urunler = self.db_yukle(self.urunler_db_path)
            self.df_satislar = self.db_yukle(self.satislar_db_path)
            self.df_stok_log = self.db_yukle(self.stok_log_db_path)
            self.df_borclar = self.db_yukle(self.borclar_db_path)
            self.df_borc_odemeler = self.db_yukle(self.borc_odemeler_db_path)
            self.df_musteriler = self.db_yukle(self.musteriler_db_path)
            self.df_firmalar = self.db_yukle(self.firmalar_db_path)
            messagebox.showinfo("Başarılı", "Tüm veritabanı başarıyla sıfırlandı.")
            self.create_main_menu()
            
    def db_disa_aktar(self):
        hedef_klasor = filedialog.askdirectory(title="Veritabanı dosyalarının kaydedileceği klasörü seçin")
        if hedef_klasor:
            try:
                shutil.copy(self.urunler_db_path, os.path.join(hedef_klasor, "urunler.csv"))
                shutil.copy(self.satislar_db_path, os.path.join(hedef_klasor, "satislar.csv"))
                shutil.copy(self.stok_log_db_path, os.path.join(hedef_klasor, "stok_log.csv"))
                shutil.copy(self.borclar_db_path, os.path.join(hedef_klasor, "borclar.csv"))
                shutil.copy(self.borc_odemeler_db_path, os.path.join(hedef_klasor, "borc_odemeler.csv"))
                shutil.copy(self.musteriler_db_path, os.path.join(hedef_klasor, "musteriler.csv"))
                shutil.copy(self.firmalar_db_path, os.path.join(hedef_klasor, "firmalar.csv"))
                messagebox.showinfo("Başarılı", f"Veritabanı dosyaları '{hedef_klasor}' klasörüne aktarıldı.")
            except Exception as e:
                messagebox.showerror("Hata", f"Dışa aktarma sırasında bir hata oluştu: {e}")

    def quit_app(self):
        self.destroy()

if __name__ == "__main__":
    app = StokYonetimProgrami()
    app.mainloop()